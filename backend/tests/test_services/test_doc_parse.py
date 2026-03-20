import pytest
from app.services.doc_parse_service import DocParseService


def test_split_into_chunks_basic():
    service = DocParseService()
    text = "A" * 500
    chunks = service._split_into_chunks(text, chunk_size=200, overlap=50)
    assert len(chunks) > 1
    assert all(len(c) <= 200 for c in chunks)


def test_split_into_chunks_short_text():
    service = DocParseService()
    text = "Short text"
    chunks = service._split_into_chunks(text, chunk_size=200)
    assert len(chunks) == 1
    assert chunks[0] == "Short text"


def test_split_into_chunks_empty():
    service = DocParseService()
    chunks = service._split_into_chunks("", chunk_size=200)
    assert chunks == []


def test_parse_unsupported_raises():
    service = DocParseService()
    with pytest.raises(ValueError, match="Unsupported"):
        service.parse(b"content", "xyz")


def test_parse_docx_bytes():
    """Test DOCX parsing with a minimal valid docx."""
    from docx import Document
    import io
    doc = Document()
    doc.add_paragraph("这是一段测试文本。")
    doc.add_paragraph("第二段内容。")
    buf = io.BytesIO()
    doc.save(buf)
    content = buf.getvalue()

    service = DocParseService()
    result = service.parse(content, "docx")
    assert "这是一段测试文本" in result
    assert "第二段内容" in result
